from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import tempfile
import os
from matplotlib.font_manager import FontProperties
def get_chinese_font():
    custom_font_path = "font/NotoSerifTC-VariableFont_wght.ttf"
    if os.path.exists(custom_font_path):
        return FontProperties(fname=custom_font_path)
    return None

ch_font = get_chinese_font()


def generate_codebook(df, column_types, variable_names, category_definitions, code_df=None, output_path="codebook.docx", preview_mode=False):
    if output_path is None:
        output_path = "codebook.docx"

    if code_df is not None:
        code_df.columns = code_df.columns.str.strip().str.lower()

    doc = Document()
    doc.add_heading("Codebook Summary Report", level=1)

    # ✅ 只統計實際存在欄位的缺失值
    valid_cols = [col for col in column_types.keys() if col in df.columns]
    
    na_counts = df[valid_cols].isnull().sum()
    na_percent = df[valid_cols].isnull().mean() * 100

    doc.add_heading("Missing Value Summary", level=2)
    na_df = pd.DataFrame({
        "column": na_counts.index,
        "missing_count": na_counts.values,
        "missing_rate (%)": na_percent.round(2).values
    }).query("`missing_count` > 0").reset_index(drop=True)

    if not na_df.empty:
        table = doc.add_table(rows=1 + len(na_df), cols=4)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Index"
        table.cell(0, 1).text = "Variable"
        table.cell(0, 2).text = "Missing Count"
        table.cell(0, 3).text = "Missing Rate (%)"
        for i, row in na_df.iterrows():
            col_name = row["column"]
            index_label = variable_names.get(col_name, col_name)
            table.cell(i + 1, 0).text = index_label
            table.cell(i + 1, 1).text = str(row["column"])
            table.cell(i + 1, 2).text = str(row["missing_count"])
            table.cell(i + 1, 3).text = str(row["missing_rate (%)"])
    else:
        doc.add_paragraph("No missing values in any columns.")

    # 🔹 變數類型統計區塊
    doc.add_heading("Variable Type Summary", level=2)
    type_count = pd.Series(column_types).value_counts().sort_index()
    type_label_map = {1: "數值型 (Numerical)", 2: "類別型 (Categorical)"}

    table = doc.add_table(rows=1 + len(type_count), cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "變數類型"
    table.cell(0, 1).text = "欄位數"
    for i, (type_code, count) in enumerate(type_count.items()):
        label = type_label_map.get(type_code, f"其他 ({type_code})")
        table.cell(i + 1, 0).text = label
        table.cell(i + 1, 1).text = str(count)

    # 🔹 欄位細節處理
    columns = code_df["variable"] if code_df is not None and "variable" in code_df.columns else df.columns

    for col in columns:
        col = str(col).strip()
        if col not in column_types or col not in df.columns:
            continue

        type_code = column_types[col]
        if type_code == 0:
            continue

        var_name = variable_names.get(col, col)
        doc.add_heading(f"Variable: {col} ({var_name})", level=2)

        # ➕ 加入 Description 段落
        description = None
        if code_df is not None:
            desc_col_candidates = ['description', 'desc', '說明']
            for desc_col in desc_col_candidates:
                if desc_col in code_df.columns:
                    row_match = code_df[code_df["variable"] == col]
                    if not row_match.empty:
                        description = str(row_match.iloc[0][desc_col])
                    break


        # 🟦 類別型
        if type_code == 2:
            value_counts = df[col].value_counts(dropna=False).sort_index()
            total = len(df)
            valid_count = df[col].notna().sum()
            missing_index = df[df[col].isna()].index.tolist()
            missing_count = df[col].isna().sum()
            defs = category_definitions.get(col, {})
            lines = [
                f"{int(k) if isinstance(k, float) and k.is_integer() else k}: {defs.get(k, '')} → {v} ({v/total:.2%})"
                for k, v in value_counts.items()
            ]
            summary_text = "\n".join(lines)

            table = doc.add_table(rows=6, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Variable Name"
            table.cell(0, 1).text = f"{col} ({var_name})"
            table.cell(1, 0).text = "Categories Summary"
            table.cell(1, 1).text = summary_text
            table.cell(2, 0).text = "Valid count"
            table.cell(2, 1).text = str(valid_count)
            table.cell(3, 0).text = "NoV count"
            table.cell(3, 1).text = str(missing_count)
            table.cell(4, 0).text = "NoV index"

            if missing_index:
                preview = ", ".join(map(str, missing_index[:5]))
                suffix = " ..." if len(missing_index) > 5 else ""
                table.cell(4, 1).text = preview + suffix
            else:
                table.cell(4, 1).text = "None"
            table.cell(5, 0).text = "Description"
            table.cell(5, 1).text = description if description else "No description available"
            
            fig, ax = plt.subplots()
            value_counts.plot(kind="bar", color="cornflowerblue", ax=ax)
            ax.set_title(f"Count Plot of {col}",fontproperties=ch_font)
            ax.set_ylabel("Frequency", fontproperties=ch_font)  # ✅ Y 軸標題
            for label in ax.get_yticklabels():                  # ✅ Y 軸數字字體
                label.set_fontproperties(ch_font)
            
            # ➤ 設定 x 軸標籤為字串（避免顯示 1.0, 2.0）
            ax.set_xticks(range(len(value_counts)))
            ax.set_xticklabels([
                str(int(cat)) if isinstance(cat, float) and cat.is_integer() else str(cat)
                for cat in value_counts.index
            ],fontproperties=ch_font)
            ax.set_xlabel(col,fontproperties=ch_font)
            # ➤ 在每根長條上標出數值（轉為 int 顯示）
            for i, (_, count) in enumerate(value_counts.items()):
                ax.text(i, count + 0.5, str(int(count)), ha='center', va='bottom', fontsize=8, fontproperties=ch_font)


            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp.name)
            plt.close()
            doc.add_picture(tmp.name, width=Inches(4.5))
            try: os.unlink(tmp.name)
            except PermissionError: pass

        # 🟩 數值型
        elif type_code == 1:
            try:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            except Exception:
                continue
            if df[col].dropna().empty:
                continue
            data = df[col].dropna()
            desc = data.describe()
            valid_count = len(data)
            missing_index = df[df[col].isna()].index.tolist()
            missing_count = len(missing_index)
            table = doc.add_table(rows=8, cols=4)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Index"
            table.cell(0, 1).text = var_name
            table.cell(0, 2).text = "Variable Name"
            table.cell(0, 3).text = col

            table.cell(1, 0).text = "Mean"
            table.cell(1, 1).text = f"{desc['mean']:.3f}"
            table.cell(1, 2).text = "Std Dev"
            table.cell(1, 3).text = f"{desc['std']:.3f}"

            table.cell(2, 0).text = "Max"
            table.cell(2, 1).text = f"{desc['max']:.3f}"
            table.cell(2, 2).text = "Min"
            table.cell(2, 3).text = f"{desc['min']:.3f}"

            table.cell(3, 0).text = "Q1 (25%)"
            table.cell(3, 1).text = f"{desc['25%']:.3f}"
            table.cell(3, 2).text = "Q2 (50%)"
            table.cell(3, 3).text = f"{desc['50%']:.3f}"

            table.cell(4, 0).text = "Q3 (75%)"
            table.cell(4, 1).text = f"{desc['75%']:.3f}"
            table.cell(4, 2).text = "Range"
            table.cell(4, 3).text = f"{desc['max'] - desc['min']:.3f}"
            
            table.cell(5, 0).text = "Valid N"
            table.cell(5, 1).text = str(valid_count)
            table.cell(5, 2).text = "Missing Count"
            table.cell(5, 3).text = str(missing_count)

            table.cell(6, 0).text = " "
            table.cell(6, 1).text = " "
            table.cell(6, 2).text = "Missing Index"
            if missing_index:
                preview = ", ".join(map(str, missing_index[:5]))
                suffix = " ..." if len(missing_index) > 5 else ""
                table.cell(6, 3).text = preview + suffix
            else:
                table.cell(6, 3).text = "None"

            table.cell(7, 0).text = "Description"
            table.cell(7, 1).merge(table.cell(7, 3))  # 合併單元格
            table.cell(7, 1).text = description if description else "No description available"

            q1 = desc['25%']
            q2 = desc['50%']
            q3 = desc['75%']
            minimum = desc['min']
            maximum = desc['max']

            # ➤ 畫圖
            fig2, ax2 = plt.subplots()
            box = ax2.boxplot([data], vert=True, patch_artist=True,
                            boxprops=dict(facecolor='lightblue', color='black'),
                            medianprops=dict(color='red'))

            ax2.set_title(f"Boxplot of {col}",fontproperties=ch_font)
            for label in ax2.get_yticklabels():
                label.set_fontproperties(ch_font)

            ax2.set_xticks([1])
            ax2.set_xticklabels([col],fontproperties=ch_font)
            ax2.set_ylabel("Value",fontproperties=ch_font)
            # ➤ 加上數值註解
            def annotate(y, label):
                ax2.text(1.1, y, f"{label}: {y:.2f}", va="center", fontsize=8, fontproperties=ch_font)

            annotate(minimum, "Min")
            annotate(q1, "Q1")
            annotate(q2, "Median")
            annotate(q3, "Q3")
            annotate(maximum, "Max")

            tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp2.name)
            plt.close()
            doc.add_picture(tmp2.name, width=Inches(4.5))
            try: os.unlink(tmp2.name)
            except PermissionError: pass

            # ➤ 畫 histogram
            fig3, ax3 = plt.subplots()
            ax3.hist(data, bins='auto', color='lightblue', edgecolor='black')
            ax3.set_title(f"Histogram of {col}",fontproperties=ch_font)
            ax3.set_xlabel(col,fontproperties=ch_font)
            ax3.set_ylabel("Frequency",fontproperties=ch_font)
            for label in ax3.get_yticklabels():
                label.set_fontproperties(ch_font)

            tmp3 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.tight_layout()
            plt.savefig(tmp3.name)
            plt.close()
            doc.add_picture(tmp3.name, width=Inches(4.5))
            try: os.unlink(tmp3.name)
            except PermissionError: pass
    doc.save(output_path)
    return output_path
