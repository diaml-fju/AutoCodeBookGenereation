from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import seaborn as sns

def generate_codebook_fast(
    df, column_types, variable_names, category_definitions, 
    code_df=None, output_path="codebook_fast.docx", 
    include_figures=True, include_kde=False  # ← 新增 KDE 選配
):
    if output_path is None:
        output_path = "codebook_fast.docx"

    doc = Document()
    doc.add_heading("Codebook Summary Report (Fast Mode)", level=1)

    # ✅ 缺失值統計
    valid_cols = [col for col in column_types.keys() if col in df.columns]
    na_counts = df[valid_cols].isnull().sum()
    na_percent = df[valid_cols].isnull().mean() * 100
    na_df = pd.DataFrame({
        "column": na_counts.index,
        "missing_count": na_counts.values,
        "missing_rate (%)": na_percent.round(2).values
    }).query("`missing_count` > 0").reset_index(drop=True)

    doc.add_heading("Missing Value Summary", level=2)
    if not na_df.empty:
        table = doc.add_table(rows=1 + len(na_df), cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Variable"
        table.cell(0, 1).text = "Missing Count"
        table.cell(0, 2).text = "Missing Rate (%)"
        for i, row in na_df.iterrows():
            table.cell(i+1, 0).text = str(row["column"])
            table.cell(i+1, 1).text = str(row["missing_count"])
            table.cell(i+1, 2).text = str(row["missing_rate (%)"])
    else:
        doc.add_paragraph("No missing values in any columns.")

    # 🔹 變數細節
    columns = code_df["variable"] if code_df is not None and "variable" in code_df.columns else df.columns
    for col in columns:
        col = str(col).strip()
        if col not in column_types or col not in df.columns:
            continue

        var_name = variable_names.get(col, col)
        doc.add_heading(f"Variable: {col} ({var_name})", level=2)

        # 數值型
        if column_types[col] == 1:
            data = pd.to_numeric(df[col], errors="coerce").dropna()
            if data.empty:
                continue
            desc = data.describe()

            table = doc.add_table(rows=5, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Mean";      table.cell(0, 1).text = f"{desc['mean']:.3f}"
            table.cell(1, 0).text = "Std Dev";   table.cell(1, 1).text = f"{desc['std']:.3f}"
            table.cell(2, 0).text = "Min";       table.cell(2, 1).text = f"{desc['min']:.3f}"
            table.cell(3, 0).text = "Max";       table.cell(3, 1).text = f"{desc['max']:.3f}"
            table.cell(4, 0).text = "Count";     table.cell(4, 1).text = str(len(data))

            if include_figures:
                # Boxplot
                buf = BytesIO()
                fig, ax = plt.subplots()
                ax.boxplot([data], vert=True, patch_artist=True,
                           boxprops=dict(facecolor='lightblue', color='black'),
                           medianprops=dict(color='red'))
                ax.set_title(f"Boxplot of {col}")
                ax.set_xticks([1]); ax.set_xticklabels([col])
                plt.tight_layout(); plt.savefig(buf, format="png", dpi=72); plt.close(fig)
                buf.seek(0); doc.add_picture(buf, width=Inches(4.0)); buf.close()

                # Histogram
                buf = BytesIO()
                fig, ax = plt.subplots()
                if np.allclose(data, data.astype(int)):  # 整數型 → 每個整數一格
                    bins = np.arange(data.min(), data.max() + 2) - 0.5
                else:
                    bins = "auto"
                ax.hist(data, bins=bins, color='lightblue', edgecolor='black')
                ax.set_title(f"Histogram of {col}")
                ax.set_xlabel(col); ax.set_ylabel("Frequency")
                plt.tight_layout(); plt.savefig(buf, format="png", dpi=72); plt.close(fig)
                buf.seek(0); doc.add_picture(buf, width=Inches(4.0)); buf.close()


        # 類別型
        elif column_types[col] == 2:
            value_counts = df[col].value_counts(dropna=False)
            total = len(df)

            summary_text = "\n".join([
                f"{k}: {v} ({v/total:.1%})"
                for k, v in value_counts.items()
            ])

            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Summary"
            table.cell(0, 1).text = summary_text
            table.cell(1, 0).text = "Count"
            table.cell(1, 1).text = str(total)

            if include_figures:
                buf = BytesIO()
                fig, ax = plt.subplots()
                value_counts.plot(kind="bar", color="cornflowerblue", ax=ax)
                ax.set_title(f"Count Plot of {col}")
                ax.set_xlabel(col); ax.set_ylabel("Frequency")
                plt.tight_layout(); plt.savefig(buf, format="png", dpi=72); plt.close(fig)
                buf.seek(0); doc.add_picture(buf, width=Inches(4.0)); buf.close()

    doc.save(output_path)
    return output_path
